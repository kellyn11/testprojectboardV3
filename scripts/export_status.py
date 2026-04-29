import json
import os
import re
import subprocess
from pathlib import Path

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.units import mm

OWNER = os.environ["OWNER"]
REPO = os.environ["REPO"]
PROJECT_OWNER = os.environ["PROJECT_OWNER"]
PROJECT_NUMBER = int(os.environ["PROJECT_NUMBER"])

DOCX_PATH = Path("input/stories.docx")

USER_PDF_OUTPUT = Path("output/user_report.pdf")
OFFICER_PDF_OUTPUT = Path("output/project_officer_report.pdf")


def run_gh(args):
    result = subprocess.run(
        ["gh"] + args,
        capture_output=True,
        text=True,
        check=True,
    )
    return result.stdout.strip()


def clean_text(text):
    return re.sub(r"\s+", " ", (text or "").strip())


def generate_title_from_story(story):
    match = re.search(r"I want to (.*?)(?: so that|$)", story, re.IGNORECASE)
    title = match.group(1).strip() if match else story.strip()

    if not title:
        return "Untitled Story"

    return title[:1].upper() + title[1:]


def extract_story_rows_from_docx():
    doc = Document(str(DOCX_PATH))
    rows_out = []

    current_section = "General"

    for p in doc.paragraphs:
        txt = clean_text(p.text)
        if re.match(r"^\d+\.\d+\s+", txt):
            current_section = txt

    for table in doc.tables:
        rows = table.rows
        if not rows:
            continue

        header = [clean_text(c.text) for c in rows[0].cells[:2]]

        if header != ["User Stories", "Acceptance Criteria"]:
            continue

        for row in rows[1:]:
            cells = row.cells
            if len(cells) < 2:
                continue

            story = clean_text(cells[0].text)
            ac_raw = cells[1].text.strip()

            if not story:
                continue

            ac_lines = []
            for line in ac_raw.splitlines():
                line = clean_text(line)
                if line:
                    ac_lines.append(line)

            rows_out.append({
                "title": generate_title_from_story(story),
                "story": story,
                "acceptance": ac_lines,
                "section": current_section,
            })

    return rows_out


def get_project_status_map():
    query = """
    query($owner:String!, $number:Int!) {
      user(login:$owner) {
        projectV2(number:$number) {
          items(first:100) {
            nodes {
              content {
                ... on Issue {
                  number
                  title
                }
              }
              fieldValues(first:20) {
                nodes {
                  ... on ProjectV2ItemFieldSingleSelectValue {
                    name
                    field {
                      ... on ProjectV2SingleSelectField {
                        name
                      }
                    }
                  }
                }
              }
            }
          }
        }
      }
    }
    """

    output = run_gh([
        "api", "graphql",
        "-f", f"query={query}",
        "-F", f"owner={PROJECT_OWNER}",
        "-F", f"number={PROJECT_NUMBER}",
    ])

    data = json.loads(output)
    items = data["data"]["user"]["projectV2"]["items"]["nodes"]

    status_map = {}

    for item in items:
        content = item.get("content")
        if not content:
            continue

        issue_title = clean_text(content.get("title", ""))
        issue_number = content.get("number")

        status = "No Status"

        for fv in item.get("fieldValues", {}).get("nodes", []):
            field = fv.get("field")
            if field and field.get("name") == "Status":
                status = fv.get("name", "No Status")
                break

        status_map[issue_title] = {
            "status": status,
            "issue_number": issue_number,
        }

    return status_map


def marker_for_status(status):
    s = (status or "").strip().lower()

    if s == "done":
        return "[X]"
    if s in ["in progress", "qa review"]:
        return "[-]"
    return "[ ]"


def build_table_data(rows, status_map, persona):
    if persona == "user":
        table_data = [[
            Paragraph("<b>Status</b>", styles["BodyText"]),
            Paragraph("<b>Story</b>", styles["BodyText"]),
            Paragraph("<b>Acceptance Criteria</b>", styles["BodyText"]),
        ]]

        for row in rows:
            info = status_map.get(row["title"], {})
            status = info.get("status", "No Status")
            marker = marker_for_status(status)

            ac_text = "<br/>".join(row["acceptance"]) if row["acceptance"] else ""

            table_data.append([
                Paragraph(f"{marker}<br/>{status}", styles["BodyText"]),
                Paragraph(row["title"], styles["BodyText"]),
                Paragraph(ac_text, styles["BodyText"]),
            ])

    else:
        table_data = [[
            Paragraph("<b>Issue #</b>", styles["BodyText"]),
            Paragraph("<b>Status</b>", styles["BodyText"]),
            Paragraph("<b>Section</b>", styles["BodyText"]),
            Paragraph("<b>Story</b>", styles["BodyText"]),
            Paragraph("<b>Acceptance Criteria</b>", styles["BodyText"]),
        ]]

        for row in rows:
            info = status_map.get(row["title"], {})
            status = info.get("status", "No Status")
            issue_number = info.get("issue_number", "-")
            marker = marker_for_status(status)

            ac_text = "<br/>".join(row["acceptance"]) if row["acceptance"] else ""

            table_data.append([
                Paragraph(str(issue_number), styles["BodyText"]),
                Paragraph(f"{marker}<br/>{status}", styles["BodyText"]),
                Paragraph(row["section"], styles["BodyText"]),
                Paragraph(row["title"], styles["BodyText"]),
                Paragraph(ac_text, styles["BodyText"]),
            ])

    return table_data


def write_pdf_report(rows, status_map, output_path, persona):
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=landscape(A4),
        leftMargin=10 * mm,
        rightMargin=10 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
    )

    elements = []

    if persona == "user":
        title = "User Story Status Report"
    else:
        title = "Project Officer Status Report"

    elements.append(Paragraph(title, styles["Title"]))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(
        "Legend: [X] Done &nbsp;&nbsp;&nbsp; [-] In Progress / QA Review &nbsp;&nbsp;&nbsp; [ ] Backlog / Todo / No Status",
        styles["BodyText"]
    ))
    elements.append(Spacer(1, 10))

    table_data = build_table_data(rows, status_map, persona)

    if persona == "user":
        col_widths = [35 * mm, 95 * mm, 140 * mm]
    else:
        col_widths = [20 * mm, 35 * mm, 45 * mm, 80 * mm, 95 * mm]

    table = Table(table_data, colWidths=col_widths, repeatRows=1)

    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))

    elements.append(table)

    done_count = sum(
        1 for row in rows
        if status_map.get(row["title"], {}).get("status") == "Done"
    )
    total_count = len(rows)
    progress = round((done_count / total_count) * 100) if total_count else 0

    elements.append(Spacer(1, 10))
    elements.append(Paragraph(f"<b>Completion:</b> {done_count} / {total_count} Completed", styles["BodyText"]))
    elements.append(Paragraph(f"<b>Progress:</b> {progress}%", styles["BodyText"]))

    doc.build(elements)


def main():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(f"Missing {DOCX_PATH}")

    rows = extract_story_rows_from_docx()

    if not rows:
        raise RuntimeError("No stories found. Header must be: User Stories | Acceptance Criteria")

    status_map = get_project_status_map()

    write_pdf_report(rows, status_map, USER_PDF_OUTPUT, "user")
    write_pdf_report(rows, status_map, OFFICER_PDF_OUTPUT, "officer")

    print(f"Generated {USER_PDF_OUTPUT}")
    print(f"Generated {OFFICER_PDF_OUTPUT}")


styles = getSampleStyleSheet()
styles["BodyText"].fontName = "Helvetica"
styles["BodyText"].fontSize = 8
styles["BodyText"].leading = 10

if __name__ == "__main__":
    main()
